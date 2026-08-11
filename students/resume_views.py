import io

from django.http import HttpResponse
from rest_framework import generics, status
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.views import APIView

from users.permissions import IsStudent
from .models import ResumeDraft
from .serializers import ResumeDraftSerializer, ResumeDraftListSerializer


def _empty_resume_data():
    return {
        "contact": {"full_name": "", "email": "", "phone": "", "location": "", "linkedin": "", "github": "", "portfolio": ""},
        "target_title": "",
        "summary": "",
        "experience": [],
        "education": [],
        "skills": [],
        "projects": [],
        "certifications": [],
    }


class ResumeDraftListCreateView(generics.ListCreateAPIView):
    """
    GET  /api/students/resumes/  -> "Open Saved" list (lightweight)
    POST /api/students/resumes/  -> create a new draft (blank, or with initial `data`)
    """
    permission_classes = [IsAuthenticated, IsStudent]

    def get_serializer_class(self):
        return ResumeDraftListSerializer if self.request.method == "GET" else ResumeDraftSerializer

    def get_queryset(self):
        return ResumeDraft.objects.filter(student=self.request.user.student_profile)

    def perform_create(self, serializer):
        data = serializer.validated_data.get("data") or _empty_resume_data()
        serializer.save(student=self.request.user.student_profile, data=data)


class ResumeDraftDetailView(generics.RetrieveUpdateDestroyAPIView):
    """GET/PATCH/DELETE /api/students/resumes/<id>/ — one draft, own only."""
    serializer_class = ResumeDraftSerializer
    permission_classes = [IsAuthenticated, IsStudent]

    def get_queryset(self):
        return ResumeDraft.objects.filter(student=self.request.user.student_profile)


class ResumePrefillView(APIView):
    """
    GET /api/students/resumes/prefill/ — "Build from Profile": returns a
    resume `data` blob pre-filled from whatever the student has already
    entered in their StudentProfile, so they aren't starting from nothing.
    """
    permission_classes = [IsAuthenticated, IsStudent]

    def get(self, request):
        profile = request.user.student_profile
        data = _empty_resume_data()
        data["contact"] = {
            "full_name": profile.full_name,
            "email": profile.personal_email or profile.college_email,
            "phone": profile.phone,
            "location": profile.current_location or profile.current_residence,
            "linkedin": profile.linkedin,
            "github": profile.github,
            "portfolio": "",
        }
        data["target_title"] = profile.course
        education_entries = []
        if profile.graduation_course or profile.course:
            education_entries.append({
                "school": "St. Andrews Institute of Technology & Management",
                "degree": profile.graduation_course or profile.course,
                "location": "Gurugram, India",
                "start_date": "",
                "end_date": profile.batch,
                "percentage_gpa": str(profile.cgpa) if profile.cgpa else "",
            })
        if profile.twelfth_percentage:
            education_entries.append({
                "school": "Senior Secondary (12th)",
                "degree": f"{profile.twelfth_board} Board" if profile.twelfth_board else "",
                "location": "",
                "start_date": "",
                "end_date": profile.twelfth_year_of_passing,
                "percentage_gpa": profile.twelfth_percentage,
            })
        data["education"] = education_entries
        if profile.certifications:
            data["certifications"] = [c.strip() for c in profile.certifications.split("\n") if c.strip()]
        if profile.achievements:
            data["summary"] = profile.achievements
        return Response(data)


class ResumeExportPDFView(APIView):
    """GET /api/students/resumes/<id>/export/pdf/ — real, selectable-text PDF (not a screenshot)."""
    permission_classes = [IsAuthenticated, IsStudent]

    def get(self, request, pk):
        try:
            draft = ResumeDraft.objects.get(pk=pk, student=request.user.student_profile)
        except ResumeDraft.DoesNotExist:
            return Response({"detail": "Resume not found."}, status=status.HTTP_404_NOT_FOUND)

        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=0.6 * inch, bottomMargin=0.6 * inch,
                                 leftMargin=0.7 * inch, rightMargin=0.7 * inch)
        styles = getSampleStyleSheet()
        name_style = ParagraphStyle("Name", parent=styles["Title"], fontSize=20, alignment=TA_CENTER, spaceAfter=2)
        contact_style = ParagraphStyle("Contact", parent=styles["Normal"], fontSize=9.5, alignment=TA_CENTER, textColor="#444444", spaceAfter=10)
        section_style = ParagraphStyle("Section", parent=styles["Heading2"], fontSize=11.5, spaceBefore=10, spaceAfter=4, textColor="#1B2A4A")
        body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, leading=14)
        bullet_style = ParagraphStyle("Bullet", parent=body_style, leftIndent=14, spaceAfter=2)
        role_style = ParagraphStyle("Role", parent=styles["Normal"], fontSize=10.5, spaceBefore=6, fontName="Helvetica-Bold")
        meta_style = ParagraphStyle("Meta", parent=styles["Normal"], fontSize=9, textColor="#666666", spaceAfter=2)

        data = draft.data or {}
        contact = data.get("contact", {})
        story = []

        story.append(Paragraph(contact.get("full_name", "") or draft.student.full_name, name_style))
        contact_line = " | ".join(filter(None, [
            contact.get("email"), contact.get("phone"), contact.get("location"),
            contact.get("linkedin"), contact.get("github"), contact.get("portfolio"),
        ]))
        if contact_line:
            story.append(Paragraph(contact_line, contact_style))

        if data.get("target_title"):
            story.append(Paragraph(data["target_title"], ParagraphStyle("Target", parent=body_style, alignment=TA_CENTER, textColor="#1B2A4A", spaceAfter=8)))

        story.append(HRFlowable(width="100%", color="#1B2A4A", thickness=1))

        if data.get("summary"):
            story.append(Paragraph("PROFESSIONAL SUMMARY", section_style))
            story.append(Paragraph(data["summary"], body_style))

        if data.get("experience"):
            story.append(Paragraph("WORK EXPERIENCE", section_style))
            for exp in data["experience"]:
                header = f"{exp.get('role', '')} — {exp.get('company', '')}"
                story.append(Paragraph(header, role_style))
                dates = " to ".join(filter(None, [exp.get("start_date"), "Present" if exp.get("current") else exp.get("end_date")]))
                meta = " | ".join(filter(None, [exp.get("location"), dates]))
                if meta:
                    story.append(Paragraph(meta, meta_style))
                for bullet in exp.get("bullets", []):
                    if bullet.strip():
                        story.append(Paragraph(f"• {bullet}", bullet_style))

        if data.get("education"):
            story.append(Paragraph("EDUCATION", section_style))
            for edu in data["education"]:
                header = f"{edu.get('degree', '')} — {edu.get('school', '')}"
                story.append(Paragraph(header, role_style))
                meta = " | ".join(filter(None, [
                    edu.get("location"), edu.get("end_date"),
                    f"Score: {edu.get('percentage_gpa')}" if edu.get("percentage_gpa") else None,
                ]))
                if meta:
                    story.append(Paragraph(meta, meta_style))

        if data.get("projects"):
            story.append(Paragraph("PROJECTS", section_style))
            for proj in data["projects"]:
                story.append(Paragraph(proj.get("name", ""), role_style))
                if proj.get("description"):
                    story.append(Paragraph(proj["description"], body_style))
                if proj.get("link"):
                    story.append(Paragraph(proj["link"], meta_style))

        if data.get("skills"):
            story.append(Paragraph("SKILLS", section_style))
            story.append(Paragraph(", ".join(data["skills"]), body_style))

        if data.get("certifications"):
            story.append(Paragraph("CERTIFICATIONS", section_style))
            for cert in data["certifications"]:
                story.append(Paragraph(f"• {cert}", bullet_style))

        doc.build(story)
        buf.seek(0)
        response = HttpResponse(buf.read(), content_type="application/pdf")
        filename = (contact.get("full_name") or draft.student.full_name or "resume").replace(" ", "_")
        response["Content-Disposition"] = f'attachment; filename="{filename}_resume.pdf"'
        return response


class ResumeExportDOCXView(APIView):
    """GET /api/students/resumes/<id>/export/docx/ — real, editable .docx."""
    permission_classes = [IsAuthenticated, IsStudent]

    def get(self, request, pk):
        try:
            draft = ResumeDraft.objects.get(pk=pk, student=request.user.student_profile)
        except ResumeDraft.DoesNotExist:
            return Response({"detail": "Resume not found."}, status=status.HTTP_404_NOT_FOUND)

        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        data = draft.data or {}
        contact = data.get("contact", {})
        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)

        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_p.add_run(contact.get("full_name") or draft.student.full_name)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)

        contact_line = " | ".join(filter(None, [
            contact.get("email"), contact.get("phone"), contact.get("location"),
            contact.get("linkedin"), contact.get("github"), contact.get("portfolio"),
        ]))
        if contact_line:
            p = doc.add_paragraph(contact_line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(9.5)

        if data.get("target_title"):
            p = doc.add_paragraph(data["target_title"])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)

        def add_section(title):
            h = doc.add_heading(title, level=2)
            for r in h.runs:
                r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
                r.font.size = Pt(12.5)

        if data.get("summary"):
            add_section("Professional Summary")
            doc.add_paragraph(data["summary"])

        if data.get("experience"):
            add_section("Work Experience")
            for exp in data["experience"]:
                p = doc.add_paragraph()
                p.add_run(f"{exp.get('role', '')} — {exp.get('company', '')}").bold = True
                dates = " to ".join(filter(None, [exp.get("start_date"), "Present" if exp.get("current") else exp.get("end_date")]))
                meta = " | ".join(filter(None, [exp.get("location"), dates]))
                if meta:
                    mp = doc.add_paragraph(meta)
                    mp.runs[0].font.size = Pt(9)
                    mp.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                for bullet in exp.get("bullets", []):
                    if bullet.strip():
                        doc.add_paragraph(bullet, style="List Bullet")

        if data.get("education"):
            add_section("Education")
            for edu in data["education"]:
                p = doc.add_paragraph()
                p.add_run(f"{edu.get('degree', '')} — {edu.get('school', '')}").bold = True
                meta = " | ".join(filter(None, [
                    edu.get("location"), edu.get("end_date"),
                    f"Score: {edu.get('percentage_gpa')}" if edu.get("percentage_gpa") else None,
                ]))
                if meta:
                    mp = doc.add_paragraph(meta)
                    mp.runs[0].font.size = Pt(9)

        if data.get("projects"):
            add_section("Projects")
            for proj in data["projects"]:
                doc.add_paragraph().add_run(proj.get("name", "")).bold = True
                if proj.get("description"):
                    doc.add_paragraph(proj["description"])
                if proj.get("link"):
                    doc.add_paragraph(proj["link"])

        if data.get("skills"):
            add_section("Skills")
            doc.add_paragraph(", ".join(data["skills"]))

        if data.get("certifications"):
            add_section("Certifications")
            for cert in data["certifications"]:
                doc.add_paragraph(cert, style="List Bullet")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        response = HttpResponse(
            buf.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        filename = (contact.get("full_name") or draft.student.full_name or "resume").replace(" ", "_")
        response["Content-Disposition"] = f'attachment; filename="{filename}_resume.docx"'
        return response